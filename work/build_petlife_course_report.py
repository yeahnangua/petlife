from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BASE_DIR = Path(__file__).resolve().parents[1]
OUT_DOCX = BASE_DIR / "PetLife_课程报告_按模板填写.docx"

DIAGRAM_DIR = BASE_DIR / "work" / "original_media" / "word" / "media"
SCREEN_DIR = BASE_DIR / "work" / "screenshots"


FIGURES = {
    "architecture": DIAGRAM_DIR / "image1.png",
    "mobile_structure": DIAGRAM_DIR / "image2.png",
    "home_modules": DIAGRAM_DIR / "image3.png",
    "backend_layers": DIAGRAM_DIR / "image4.png",
    "purchase_flow": DIAGRAM_DIR / "image5.png",
    "booking_flow": DIAGRAM_DIR / "image6.png",
    "home": SCREEN_DIR / "petlife-mobile-home.png",
    "product": SCREEN_DIR / "petlife-mobile-product-detail.png",
    "cart": SCREEN_DIR / "petlife-mobile-cart.png",
    "order_confirm": SCREEN_DIR / "petlife-mobile-order-confirm.png",
    "booking_confirm": SCREEN_DIR / "petlife-mobile-booking-confirm.png",
    "admin_dashboard": SCREEN_DIR / "petlife-admin-dashboard-desktop.png",
    "admin_products": SCREEN_DIR / "petlife-admin-products-desktop.png",
}


ACCENT = RGBColor(46, 116, 181)
DARK_ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)


def set_font(run, size=None, bold=None, color=None, font="SimSun"):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font)


def set_paragraph_font(paragraph, size=11, bold=False, color=None, font="SimSun"):
    for run in paragraph.runs:
        set_font(run, size=size, bold=bold, color=color, font=font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def set_table_grid(table, widths):
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=ACCENT, font="SimHei")
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_font(run, size=13, bold=True, color=ACCENT, font="SimHei")
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        set_font(run, size=12, bold=True, color=DARK_ACCENT, font="SimHei")
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc, text, first_line=True):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if first_line:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    run = paragraph.add_run(text)
    set_font(run, size=11, font="SimSun")
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=11, font="SimSun")
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, size=11, font="SimSun")
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(9)
    run = paragraph.add_run(text)
    set_font(run, size=9, color=MUTED, font="SimSun")
    return paragraph


def usable_image_width(path, max_width):
    with Image.open(path) as image:
        width, height = image.size
    if width == 0:
        return max_width
    return max_width


def add_figure(doc, path, caption, width_inches=6.25):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(usable_image_width(path, width_inches)))
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False
    set_table_width(table)
    if widths:
        set_table_grid(table, widths)

    header_cells = table.rows[0].cells
    for idx, text in enumerate(headers):
        cell = header_cells[idx]
        set_cell_shading(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_font(run, size=10.5, bold=True, font="SimHei")

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(str(text))
            set_font(run, size=10, font="SimSun")

    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    for line_no, line in enumerate(code.splitlines()):
        if line_no:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_font(run, size=9.5, font="Menlo")
    return paragraph


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def set_section_break_type(section, value="nextPage"):
    sect_pr = section._sectPr
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.insert(0, type_el)
    type_el.set(qn("w:val"), value)


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(60)
    run = paragraph.add_run("浙江万里学院 2025/2026 学年第二学期")
    set_font(run, size=16, bold=True, font="SimHei")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(36)
    run = paragraph.add_run("《移动互联网软件高级开发技术》大作业课程报告")
    set_font(run, size=15, bold=True, font="SimHei")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("题  目")
    set_font(run, size=18, bold=True, font="SimHei")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(42)
    run = paragraph.add_run("PetLife 宠物生活馆移动端 Web 与后台管理系统设计与实现")
    set_font(run, size=18, bold=True, font="SimHei")

    fields = [
        "★姓名：____________  班级：____________  学号：____________",
        "姓名：______________  班级：____________  学号：____________",
        "姓名：______________  班级：____________  学号：____________",
        "姓名：______________  班级：____________  学号：____________",
        "完成日期：2026 年 6 月",
    ]
    for field in fields:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(field)
        set_font(run, size=13, font="SimSun")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(96)
    run = paragraph.add_run("（★为本项目的主要负责同学）")
    set_font(run, size=10.5, font="SimSun")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    run = paragraph.add_run("大数据与软件学院")
    set_font(run, size=13, bold=True, font="SimHei")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("二零二六年六月")
    set_font(run, size=13, font="SimSun")



def add_overview(doc):
    heading = add_heading(doc, "一、项目概述（500字以内）", 1)
    heading.paragraph_format.page_break_before = True
    add_body(
        doc,
        "PetLife 宠物生活馆项目围绕“商品购买 + 服务预约 + 宠物档案 + 会员权益 + 后台运营”展开，面向真实养宠场景设计移动端 Web 应用。系统采用前后端分离架构：用户端和后台管理端均使用 Vue 3、Vite、Vue Router 与 Pinia 组织页面和状态，后端使用 Node.js、Express 与 SQLite 提供统一 API、业务规则和数据持久化。当前版本已完成商品浏览、加入购物车、修改数量、结算下单、服务预约、宠物档案、地址管理、订单/预约查看，以及后台分类、商品、服务、门店、时段、订单、预约管理等功能。通过本地部署、Playwright 页面验证和生产构建验证，项目可以完成从代码启动到访问前台首页、后台管理和关键业务流程的闭环。项目亮点是把传统宠物商城和到店服务预约结合在同一移动端体验中，并通过后台电脑端管理台支撑内容运营。",
    )


def add_deployment(doc):
    add_heading(doc, "二、部署指南", 1)
    add_body(
        doc,
        "项目提交物包含前台、后台和服务端三个可独立启动的子系统。部署时需要先准备 Node.js 与 npm 环境，服务端使用 SQLite 文件数据库，不依赖额外安装的 MySQL 或 PostgreSQL。前后台开发服务器通过 Vite 代理访问后端 `/api` 和 `/uploads`，因此三个进程需要同时运行。"
    )
    add_table(
        doc,
        ["项目", "目录", "启动命令", "访问地址或说明"],
        [
            ["服务端", "server/", "npm run dev:server", "默认监听 http://127.0.0.1:8787，提供 API 与上传文件访问。"],
            ["用户前台", "根目录 src/", "npm run dev", "默认访问 http://127.0.0.1:5173，Vite 代理到后端。"],
            ["后台管理", "admin/", "npm run dev:admin", "默认访问 http://127.0.0.1:5174，管理员密钥来自 server/.env。"],
        ],
        widths=[1600, 1900, 2100, 3760],
    )
    add_heading(doc, "2.1 从代码到首页访问的闭环步骤", 2)
    for step in [
        "在项目根目录执行 `npm install`，在 `admin/` 和 `server/` 目录分别执行依赖安装，确保三个 package-lock 对应依赖完整。",
        "检查 `server/.env`，确认 `PORT=8787`、`DB_PATH=./data/petlife.sqlite`、`ADMIN_KEY=123456`、`UPLOAD_DIR=./uploads` 等配置存在。",
        "执行 `npm run dev:server` 启动后端。后端启动时会执行数据库迁移，若用户表为空会装填演示种子数据。",
        "执行 `npm run dev` 启动移动端前台，浏览器访问 `http://127.0.0.1:5173` 可看到 PetLife 首页。",
        "执行 `npm run dev:admin` 启动后台管理端，访问 `http://127.0.0.1:5174`，输入管理员密钥后进入电脑端后台概览。",
    ]:
        add_numbered(doc, step)
    add_body(doc, "本次文档生成时已实际启动三个进程，前台、后台页面均可访问，后端接口也能返回商品、服务和后台商品列表数据。")


def add_requirements(doc):
    add_heading(doc, "三、需求分析", 1)
    add_body(
        doc,
        "项目用户分为普通养宠用户和后台管理员两类。普通用户关注商品购买、服务预约、宠物资料和会员权益；管理员关注商品目录、门店服务、时段配置、订单履约和预约处理。"
    )
    add_heading(doc, "1、用户需求分析", 2)
    add_table(
        doc,
        ["用户角色", "使用场景", "对应功能"],
        [
            ["普通用户", "第一次进入平台，希望快速理解平台提供什么服务。", "首页主视觉、宠物类型切换、快捷入口、热门服务、热卖商品、会员权益入口。"],
            ["商品购买用户", "浏览宠物食品、零食、清洁和出行用品，并完成购买。", "分类筛选、商品列表、商品详情、规格选择、加入购物车、数量修改、确认订单和订单详情。"],
            ["服务预约用户", "为宠物预约洗护、美容、健康护理或寄养服务。", "服务列表、服务详情、门店选择、日期时段选择、宠物档案选择、预约确认和预约详情。"],
            ["长期养宠用户", "需要维护宠物信息，使服务预约更准确。", "宠物档案新增、编辑、头像、品种、体重、过敏项和偏好信息维护。"],
            ["管理员", "维护商品/服务资料并处理订单、预约。", "后台概览、分类、商品、服务、门店、时段、订单、预约管理。"],
        ],
        widths=[1450, 3300, 4610],
    )
    add_heading(doc, "2、系统需求分析", 2)
    add_table(
        doc,
        ["需求类型", "具体要求", "项目实现方式"],
        [
            ["性能需求", "局域网开发环境下前台页面和后台列表应能秒级响应；商品、服务和预约数据通过分页或列表接口加载。", "Vite 前端开发服务器提供快速热更新；Express 只承担 API 入口；SQLite 用轻量查询支撑课程项目规模。"],
            ["安全性需求", "后台管理需要管理员校验，用户提交订单和预约时需要参数校验，密码或密钥不应直接写在前端。", "后台 API 使用 `x-admin-key` 校验；用户接口由演示用户中间件绑定；后端 validators 统一校验必填字段、日期和状态。"],
            ["可用性需求", "移动端主要流程应适合手机单手操作，后台应适合电脑端表格管理。", "前台使用底部导航、吸底操作条和卡片式移动页面；后台使用电脑端左侧导航和表格列表。"],
            ["数据一致性需求", "订单创建要扣减库存，取消订单要回滚库存；预约要检查时段容量。", "订单服务使用数据库事务写订单项并扣减库存；预约服务在事务内检查时段容量并生成快照数据。"],
            ["可维护性需求", "前台、后台、服务端职责清晰，便于小组协作和后续二次开发。", "根目录按 `src/`、`admin/`、`server/`、`docs/` 分层，服务端按 routes/controllers/services/repositories/db 划分。"],
        ],
        widths=[1500, 3900, 3960],
    )


def add_tech_stack(doc):
    add_heading(doc, "四、技术选型", 1)
    add_body(
        doc,
        "技术选型围绕移动互联网课程要求和项目规模确定。前后台均采用 Vue 3 生态，便于组件化开发、路由组织和状态复用；服务端采用 Node.js 与 Express，能够快速建立 REST API；数据库采用 SQLite，便于随项目代码一起提交和演示。"
    )
    add_table(
        doc,
        ["层次", "技术", "选择原因及与需求的契合度"],
        [
            ["移动端前台", "Vue 3、Vite、Vue Router、Pinia", "Vue 3 适合组件化页面开发，Vite 启动和构建速度快，Router 支撑首页、分类、商品、服务、订单、我的等页面，Pinia 管理购物车和预约状态。"],
            ["后台管理端", "Vue 3、Vite、Vue Router、Pinia", "后台需要多列表和表单管理，复用同一技术栈能减少学习成本；独立 admin 目录便于与用户端分开部署。"],
            ["服务端", "Node.js、Express、dotenv、multer", "Express 适合实现 REST API；dotenv 管理端口、数据库和管理员密钥；multer 处理商品、服务和宠物图片上传。"],
            ["数据库", "SQLite、better-sqlite3", "SQLite 无需独立数据库服务，适合课程项目演示；better-sqlite3 支持同步事务，便于保证订单和预约数据一致性。"],
            ["测试与验证", "Vitest、Supertest、Playwright MCP", "Vitest/Supertest 用于接口和组件验证；Playwright MCP 用于从真实页面截取最新前台移动端和后台电脑端界面。"],
        ],
        widths=[1500, 2500, 5360],
    )


def add_design(doc):
    add_heading(doc, "五、项目设计", 1)
    add_body(
        doc,
        "系统采用三端协同设计：移动端前台负责用户体验，后台管理端负责运营维护，Node.js 后端负责业务规则、数据持久化和统一响应。前后台通过 `/api/public`、`/api/user`、`/api/admin` 三类接口访问服务端，数据最终保存在 SQLite 数据库和上传目录中。"
    )
    add_heading(doc, "1、框架设计", 2)
    add_figure(doc, FIGURES["architecture"], "图1 项目总体架构图", 6.25)
    add_body(
        doc,
        "移动端页面被统一纳入 Vue Router，底部导航覆盖首页、分类、服务、订单和我的五个主入口；商品详情、确认订单、预约确认等深层流程通过独立路由承载。"
    )
    add_figure(doc, FIGURES["mobile_structure"], "图2 移动端页面结构图", 6.25)
    add_body(
        doc,
        "首页承担品牌展示和业务分流，包含运营横幅、宠物类型切换、快捷入口、热门服务、组合推荐、热卖商品和会员权益等模块。"
    )
    add_figure(doc, FIGURES["home_modules"], "图3 首页功能模块结构图", 6.25)
    add_body(
        doc,
        "服务端分为路由层、控制器层、服务层、仓储层和数据库层。路由层只负责 URL 分发，控制器层处理请求与响应，服务层实现业务规则，仓储层封装 SQL 操作。"
    )
    add_figure(doc, FIGURES["backend_layers"], "图4 后端分层结构图", 5.4)

    add_heading(doc, "1.1 核心业务流程", 3)
    add_body(doc, "商品购买流程从首页或分类页进入商品列表，再进入商品详情选择规格，加入购物车后修改数量并结算，最终进入订单详情。")
    add_figure(doc, FIGURES["purchase_flow"], "图5 商品购买流程图", 6.25)
    add_body(doc, "服务预约流程从服务页进入服务详情，选择门店、日期和时间段后进入预约确认页，选择宠物档案并提交，最后查看预约详情。")
    add_figure(doc, FIGURES["booking_flow"], "图6 服务预约流程图", 6.25)

    add_heading(doc, "1.2 新版界面运行截图", 3)
    add_body(
        doc,
        "以下界面截图均由 Playwright MCP 从当前运行版本重新截取。前台为 430×932 的移动端视口；后台管理端为 1440×900 的电脑端视口。"
    )
    add_figure(doc, FIGURES["home"], "图7 前台首页运行截图（移动端）", 4.2)
    add_figure(doc, FIGURES["product"], "图8 商品详情运行截图（移动端）", 4.2)
    add_figure(doc, FIGURES["cart"], "图9 购物车运行截图（移动端）", 4.2)
    add_figure(doc, FIGURES["order_confirm"], "图10 确认订单运行截图（移动端）", 4.2)
    add_figure(doc, FIGURES["booking_confirm"], "图11 预约确认运行截图（移动端）", 4.2)
    add_figure(doc, FIGURES["admin_dashboard"], "图12 后台概览运行截图（电脑端）", 6.25)
    add_figure(doc, FIGURES["admin_products"], "图13 后台商品管理运行截图（电脑端）", 6.25)


def add_interfaces(doc):
    add_heading(doc, "2、关键接口", 2)
    add_body(
        doc,
        "关键接口分为前端内部模块接口和后端 REST API 两类。前端内部接口主要体现路由、状态管理和页面之间的协作；后端 API 面向前台和后台提供数据读写能力。"
    )
    add_heading(doc, "2.1 前端内部关键模块", 3)
    add_table(
        doc,
        ["模块", "关键接口或方法", "作用"],
        [
            ["src/router/index.js", "`/`、`/category`、`/product/:id`、`/cart`、`/order/confirm`、`/service/:id`、`/booking/confirm`", "统一管理移动端页面入口，使商品购买和服务预约流程可以按路由串联。"],
            ["src/stores/cart.js", "`fetchCart()`、`addProduct(product, specLabel, quantity, specKey)`、`updateQuantity(id, delta)`、`removeItem(id)`", "管理购物车拉取、添加、选择、数量修改、删除和失效商品清理。"],
            ["src/stores/booking.js", "`prepareFromService(service, selection)`、`fetchSlots()`、`submitBooking()`、`cancelBooking(id)`", "管理服务预约的服务、门店、日期、时段、宠物档案和提交状态。"],
            ["admin/src/router/index.js", "`/login`、`/products`、`/services`、`/stores`、`/time-slots`、`/orders`、`/bookings`", "管理后台路由和登录守卫，未登录时重定向到管理员登录页。"],
        ],
        widths=[2100, 3950, 3310],
    )
    add_heading(doc, "2.2 后端对外关键 API", 3)
    add_table(
        doc,
        ["接口", "方法", "参数/请求体", "返回或作用"],
        [
            ["/api/public/categories", "GET", "可选 `petType`", "返回启用分类列表，供首页和分类页展示。"],
            ["/api/public/products", "GET", "`petType`、`categoryId`、`keyword`", "返回商品列表，包含价格、库存、标签和封面。"],
            ["/api/public/products/:id", "GET", "`id`", "返回商品详情、规格、图集、摘要和推荐信息。"],
            ["/api/user/cart", "GET", "演示用户由中间件绑定", "返回当前用户购物车列表和选中状态。"],
            ["/api/user/cart/items", "POST", "`product_id`、`spec_key`、`spec_label`、`quantity`", "新增购物车项；若商品规格相同则按后端逻辑合并或更新。"],
            ["/api/user/cart/items/:id", "PUT/DELETE", "`quantity`、`selected` 或 `id`", "修改购物车数量/选中状态或移除指定商品。"],
            ["/api/user/orders", "POST", "`address_id`、`remark`", "根据已选购物车项创建订单，事务化写入订单项并扣减库存。"],
            ["/api/user/bookings", "POST", "`pet_id`、`service_id`、`store_id`、`time_slot_id`、`booking_date`、`contact_phone`", "创建服务预约，校验宠物、服务、门店、时段和容量。"],
            ["/api/admin/products", "GET/POST", "Header: `x-admin-key`", "后台商品列表读取和新增。"],
            ["/api/admin/products/:id", "PUT/DELETE", "Header: `x-admin-key`", "后台商品编辑、上下架或删除。"],
            ["/api/admin/orders/:id/status", "POST", "`status`", "后台更新订单履约状态。"],
            ["/api/admin/bookings/:id/status", "POST", "`status`", "后台更新预约服务状态。"],
        ],
        widths=[2500, 900, 3300, 2660],
    )
    add_heading(doc, "2.3 关键业务服务签名", 3)
    add_code_block(
        doc,
        "createUserOrder(db, userId, payload) -> OrderDetail\n"
        "createUserBooking(db, userId, payload) -> Booking\n"
        "getUserCart(db, userId) -> Cart\n"
        "listAdminProducts(db, query) -> Product[]\n"
        "updateAdminOrder(db, id, payload) -> Order\n"
        "updateAdminBooking(db, id, payload) -> Booking",
    )
    add_body(
        doc,
        "订单服务的核心点是事务处理：读取已选购物车项、生成订单和订单项、扣减库存、删除已结算购物车项。预约服务的核心点是容量校验：提交前检查宠物、服务、门店和时段是否有效，并确认该时段未超过可预约容量。"
    )


def add_tests(doc):
    add_heading(doc, "六、系统测试", 1)
    add_body(
        doc,
        "系统测试围绕用户最常用路径设计，包括首页访问、商品购物车、确认订单、服务预约和后台管理。测试结果来自本次实际运行、Playwright 页面操作、后端接口读取和 Vite 构建验证。"
    )
    add_table(
        doc,
        ["功能点", "预期行为", "边界条件", "实际行为"],
        [
            ["前台首页访问", "访问 `http://127.0.0.1:5173/#/` 后显示品牌主视觉、快捷入口、热门服务和底部导航。", "后端接口可用；图片资源可加载。", "Playwright 截图显示首页正常渲染，页面标题为 PetLife 宠物生活馆。"],
            ["商品详情与购物车", "进入商品详情页后能查看规格、价格、评分和购物车入口；购物车页能显示商品、数量按钮、已选件数和结算按钮。", "库存为 0 的商品应显示售罄或无法结算；数量不能小于 1。", "实测商品详情、购物车和确认订单页面均正常展示；购物车提供减少/增加数量和去结算入口。"],
            ["确认订单", "从购物车点击去结算后进入确认订单页，展示地址、商品明细、金额和提交订单入口。", "未选择有效商品时应禁止结算；缺少地址时应提示补全地址。", "实测从购物车进入 `#/order/confirm` 成功，页面可用于下单确认。"],
            ["服务预约", "进入服务详情选择门店、日期、时段后点击立即预约，进入预约确认页并显示宠物、门店和价格。", "服务与宠物类型不匹配、时段容量已满时应阻止提交。", "实测服务详情到预约确认页流程正常，预约确认截图显示门店、日期、宠物选择和提交预约按钮。"],
            ["后台登录与商品管理", "访问后台后输入管理员密钥进入概览，商品页显示商品列表、筛选和新增/编辑入口。", "未带 `x-admin-key` 的后台 API 请求应返回未授权。", "实测后台使用电脑端 1440×900 视口登录并截图，概览与商品管理页均正常显示。"],
            ["后端公开接口", "`GET /api/public/products` 和 `GET /api/public/services` 返回 `code=0` 的列表数据。", "接口路径错误时应返回 not found。", "实测商品、服务、后台商品接口均返回 `code=0`，错误路径 `/api/public/home` 返回 not found。"],
            ["构建验证", "前台和后台能完成生产构建，说明依赖和模块引用完整。", "构建阶段若路由、组件或导入错误应失败。", "本次执行 `npm run build` 和 `npm run build:admin` 均通过，前台构建约 660ms，后台构建约 494ms。"],
        ],
        widths=[1600, 3000, 2300, 2460],
    )


def add_summary(doc):
    add_heading(doc, "七、总结与反思", 1)
    add_body(
        doc,
        "通过 PetLife 项目开发，可以更直观地理解移动互联网软件开发不是简单堆叠页面，而是要同时处理业务流程、用户体验、状态管理、接口契约和数据一致性。商品购买链路和服务预约链路表面相似，实际约束差异明显：商品侧重点在购物车、库存和订单快照；服务侧重点在宠物档案、门店时段和容量校验。如果没有清晰的状态边界，两条链路很容易互相影响。"
    )
    add_body(
        doc,
        "本项目的优点在于前台、后台和服务端边界较清楚，目录结构便于小组分工；移动端组件如商品卡片、服务卡片、吸底操作条、宠物卡片等提升了页面复用性；后端按 routes、controllers、services、repositories 分层，关键业务逻辑可以在服务层集中维护。后台管理端采用电脑端布局，适合运营人员长期处理表格和表单数据。"
    )
    add_body(
        doc,
        "开发过程中也暴露出一些问题。第一，需求持续扩展后，测试用例需要及时跟随业务字段变化，否则容易出现旧断言和新接口不一致的情况。第二，图片上传、商品状态、预约容量等边界场景需要更多自动化测试覆盖。第三，演示用户中间件适合课程项目，但真实上线时应替换为正式登录注册、密码加密、Token 鉴权和权限分级。"
    )
    add_body(
        doc,
        "后续改进方向包括：补齐完整的登录注册和权限体系；扩展购物车清空、批量删除、优惠券和支付状态；为后台增加更明确的表单校验和操作确认；为订单创建、取消、预约容量等关键路径增加更稳定的单元测试和接口测试。总体来看，项目已经覆盖宠物网站所要求的购物车模块，并在此基础上增加服务预约和后台运营能力，能够体现移动端 Web 项目从需求分析到部署验证的完整过程。"
    )


def build_document():
    missing = [str(path) for path in FIGURES.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing figure files:\n" + "\n".join(missing))

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_overview(doc)
    add_deployment(doc)
    add_requirements(doc)
    add_tech_stack(doc)
    add_design(doc)
    add_interfaces(doc)
    add_tests(doc)
    add_summary(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_document()
